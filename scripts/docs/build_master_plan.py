from __future__ import annotations

from collections.abc import Iterable
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Flowable,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.platypus import (
    Image as RLImage,
)

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "generated"
TMP = ROOT / "tmp" / "docs" / "master_plan"
DOCX_OUT = OUT / "XNative_Master_Architecture_And_Implementation_Plan.docx"
PDF_OUT = OUT / "XNative_Master_Architecture_And_Implementation_Plan.pdf"

NAVY = "17324D"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
GREEN = "26734D"
GOLD = "8A6116"
RED = "9B1C1C"
WHITE = "FFFFFF"
BLACK = "111111"


PHASES = [
    (
        "0",
        "Gercek durum, paket ve yonetisim",
        "Yaniltici tamamlanma iddialarini kaldir; tekrar uretilebilir gelistirme tabani kur.",
        "Temiz kurulum, aktif kod yolu ve kanitli belge durumu.",
    ),
    (
        "1",
        "Alan modeli ve kalici veri omurgasi",
        "Versionlanmis sozlesmeler, migration, repository, audit ve idempotency.",
        "Tekil post, tutarli foreign key ve izlenebilir feature.",
    ),
    (
        "2",
        "Guvenilir gorunur tarayici yakalama",
        "Extension, local outbox, POST /capture, retry ve DOM kalite kontrolu.",
        "Extension -> API -> DB veri kaybi olmadan calisir.",
    ),
    (
        "3",
        "Ingestion, is kuyrugu ve pipeline",
        "SQLite job queue, stage sinirlari, retry/dead-letter ve backfill.",
        "Worker restart cift yazim veya is kaybi uretmez.",
    ),
    (
        "4",
        "Medya saklama ve arsiv",
        "Exact hash, gercek pHash, content-addressed store, kota ve retention.",
        "Benzer gorsel cluster; kontrollu disk buyumesi.",
    ),
    (
        "5",
        "Multimodal icerik anlama",
        "NLP, OCR, CV, ASR, text-image iliski ve confidence-aware fusion.",
        "Metin veya medya tek basina/birlikte baglam kaniti sunar.",
    ),
    (
        "6",
        "Olay bellegi ve semantik retrieval",
        "Entity/event graph, BM25+dense retrieval, novelty ve fatigue.",
        "Kelime eslesmesi olmadan baglamsal benzerlik bulunur.",
    ),
    (
        "7",
        "Kaynak ve takip adayi ogrenimi",
        "Erken sinyal, dogruluk, risk, diversity ve drift.",
        "Adaylar kanit ve belirsizlikle sunulur; auto-follow yok.",
    ),
    (
        "8",
        "Cok hedefli siralama",
        "Composable source/hydrator/filter/scorer/selector ve multi-action utility.",
        "Aciklanabilir skor, offline metrik ve diversity kapisi.",
    ),
    (
        "9",
        "X-native taslak ve guvenlik",
        "Kosullu varyant, style retrieval, claim/risk/telif kontrolleri.",
        "Farkli olaylar farkli ve kanitli taslaklar uretir.",
    ),
    (
        "10",
        "Feedback ve kontrollu ogrenme",
        "Reward, segment baseline, bandit, drift, champion/challenger ve rollback.",
        "Yalniz yeterli ve versionlanmis kanitla model degisir.",
    ),
    (
        "11",
        "Review UI ve operasyon",
        "Inbox, suggestion queue, source review, reports, settings ve accessibility.",
        "Capture -> review -> feedback akisi UI'dan tamamlanir.",
    ),
    (
        "12",
        "Performans ve guvenlik sertlestirme",
        "SLO, metric, benchmark, localhost auth, backup ve diagnostics.",
        "SLO regression ve restore testi gecer.",
    ),
    (
        "13",
        "Test matrisi ve release",
        "Unit, integration, Docker, multimodal golden set, UI E2E ve privacy.",
        "Tum P0/P1 kabul kapilari kanitlidir.",
    ),
    (
        "14",
        "Dokumantasyon ve teslimat",
        "C4/DFD/sequence/ERD, runbook, model/data card ve senkron belgeler.",
        "Markdown, DOCX ve PDF ayni release'i anlatir.",
    ),
]

COMPLEXITY = [
    ("Capture normalize", "Alan bazli parse", "O(N)", "O(1)/post"),
    ("Exact dedup", "Indexed SHA/post ID", "O(1) ort.", "O(N) indeks"),
    ("pHash kucuk veri", "Hamming scan", "O(M)", "O(M)"),
    ("pHash buyuk veri", "Bucket/LSH", "O(B) beklenen", "O(M)"),
    ("Lexical retrieval", "SQLite FTS5/BM25", "~O(log N + K)", "Indeks"),
    ("Dense retrieval", "Dot product / HNSW", "O(Nd) / ~O(log N)", "O(Nd)"),
    ("Rule scoring", "F feature", "O(NF)", "O(F)"),
    ("MMR selection", "Top-K pairwise", "O(K^2)", "O(K^2)"),
    ("Event clustering", "Windowed ANN", "~O(N log N)", "O(N)"),
    ("Video analysis", "S sampled frame", "O(S inference)", "Sinirli"),
]

SLOS = [
    ("POST /capture", "p95 < 150 ms", "Agir analiz yok; 202 + kalici job"),
    ("Inbox query", "p95 < 300 ms", "100 bin kayit, pagination"),
    ("Cheap text feature", "p95 < 100 ms/post", "CPU baseline"),
    ("Image hash + thumbnail", "p95 < 300 ms/media", "Cache miss"),
    ("Suggestion queue", "p95 < 500 ms", "Feature hazir"),
    ("Data loss", "0", "Retry/dead-letter ve idempotency"),
]

ALGORITHM_LADDER = [
    (
        "Text",
        "TF-IDF/Hashing + logistic/SGD",
        "MultinomialNB speed baseline; small multilingual encoder",
        "LSTM only after >=50k clean sequence labels and measured uplift",
    ),
    (
        "Events",
        "Time-window graph + union-find",
        "DBSCAN/HDBSCAN benchmark",
        "Do not treat global K-Means clusters as live event identity",
    ),
    (
        "Fusion",
        "Calibrated late-fusion logistic",
        "HistGradientBoosting/LightGBM",
        "No end-to-end multimodal training on target hardware",
    ),
    (
        "Ranking",
        "Explainable multi-action utility",
        "Calibrated logistic, then pairwise ranker",
        "No single engagement objective or uncontrolled neural ranker",
    ),
    (
        "Exploration",
        "Thompson Sampling",
        "LinUCB when context is sufficient",
        "Never explore risk, claims or public actions",
    ),
    (
        "Optimization",
        "Random search / successive halving",
        "TPE for conditional search spaces",
        "No genetic/PSO scheduler without benchmark evidence",
    ),
]

SOURCES = [
    "xAI, X For You Feed Algorithm, https://github.com/xai-org/x-algorithm (accessed 2026-06-22).",
    "Groq, Rate Limits, https://console.groq.com/docs/rate-limits (accessed 2026-06-22).",
    "Hugging Face, Downloading Files, https://huggingface.co/docs/huggingface_hub/main/package_reference/file_download (accessed 2026-06-22).",
    "Google One, Google AI Plans, https://one.google.com/about/google-ai-plans/ (accessed 2026-06-22).",
    "Google AI for Developers, Gemini API Billing, https://ai.google.dev/gemini-api/docs/billing (accessed 2026-06-22).",
    "Google AI for Developers, Gemini API Rate Limits, https://ai.google.dev/gemini-api/docs/rate-limits (accessed 2026-06-22).",
    "Google AI for Developers, Gemini API Pricing, https://ai.google.dev/gemini-api/docs/pricing (accessed 2026-06-22).",
    "scikit-learn, SGDClassifier, https://scikit-learn.org/stable/modules/generated/sklearn.linear_model.SGDClassifier.html (accessed 2026-06-22).",
    "scikit-learn, Probability calibration, https://scikit-learn.org/stable/modules/calibration.html (accessed 2026-06-22).",
    "scikit-learn, Out-of-core text classification, https://scikit-learn.org/stable/auto_examples/applications/plot_out_of_core_classification.html (accessed 2026-06-22).",
    "scikit-learn, Naive Bayes, https://scikit-learn.org/stable/modules/naive_bayes.html (accessed 2026-06-22).",
    "OpenAI, CLIP, https://arxiv.org/abs/2103.00020 (accessed 2026-06-22).",
    "Meta AI Research, FAISS, https://github.com/facebookresearch/faiss/wiki (accessed 2026-06-22).",
    "Milvus, Standalone prerequisites, https://milvus.io/docs/prerequisite-docker.md (accessed 2026-06-22).",
    "El-Kishky et al., TwHIN, https://arxiv.org/abs/2202.05387 (accessed 2026-06-22).",
    "Hu et al., Heterogeneous Graph Transformer, https://arxiv.org/abs/2003.01332 (accessed 2026-06-22).",
]


def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    name = "Arial Bold.ttf" if bold else "Arial.ttf"
    return ImageFont.truetype(f"/System/Library/Fonts/Supplemental/{name}", size)


def _arrow(
    draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = "#36566F"
) -> None:
    draw.line([start, end], fill=fill, width=4)
    x, y = end
    draw.polygon([(x, y), (x - 12, y - 7), (x - 12, y + 7)], fill=fill)


def create_architecture_image(path: Path) -> None:
    width, height = 1800, 760
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title = _font(34, True)
    body = _font(22, True)
    small = _font(17)
    draw.text((60, 35), "XNative hedef mimarisi", font=title, fill="#17324D")
    boxes = [
        (60, 150, 280, 290, "Capture", "Extension\nmanual import", "#E8F1FB"),
        (350, 150, 570, 290, "Local API", "validation\ninbox", "#E8F1FB"),
        (640, 150, 860, 290, "Job queue", "retry\ndead-letter", "#F8F0E6"),
        (930, 100, 1190, 250, "Multimodal", "NLP OCR CV\nASR fusion", "#EDF7EF"),
        (930, 330, 1190, 480, "Event memory", "FTS + vector\nentity graph", "#EDF7EF"),
        (1260, 150, 1510, 290, "Rank + generate", "multi-action\ndiversity policy", "#F8F0E6"),
        (1560, 150, 1760, 290, "Review UI", "approve edit\nreject ignore", "#FBECEC"),
        (1260, 400, 1510, 540, "Learning", "evaluation\nregistry rollback", "#F4ECFA"),
        (350, 470, 700, 620, "Storage", "SQLite FTS5 audit | content-addressed media", "#F2F4F7"),
    ]
    for x1, y1, x2, y2, heading, detail, fill in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=20, fill=fill, outline="#8AA4B8", width=3)
        draw.text((x1 + 18, y1 + 18), heading, font=body, fill="#17324D")
        draw.multiline_text((x1 + 18, y1 + 60), detail, font=small, fill="#344054", spacing=7)
    _arrow(draw, (280, 220), (350, 220))
    _arrow(draw, (570, 220), (640, 220))
    _arrow(draw, (860, 220), (930, 175))
    _arrow(draw, (1060, 250), (1060, 330))
    _arrow(draw, (1190, 220), (1260, 220))
    _arrow(draw, (1510, 220), (1560, 220))
    _arrow(draw, (1660, 290), (1510, 460))
    _arrow(draw, (1385, 400), (1385, 290))
    draw.line([(460, 290), (460, 470)], fill="#36566F", width=4)
    draw.line([(750, 290), (750, 550), (700, 550)], fill="#36566F", width=4)
    draw.text(
        (60, 690),
        "Otonom: analiz, siralama, raporlama ve kontrollu ogrenme | Insan onayi: yayin ve takip",
        font=_font(21, True),
        fill="#9B1C1C",
    )
    img.save(path)


def create_learning_image(path: Path) -> None:
    width, height = 1700, 680
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 35), "Kontrollu ogrenme dongusu", font=_font(34, True), fill="#17324D")
    nodes = [
        (70, 210, "Evidence", "capture + context"),
        (330, 210, "Ranking", "cok hedefli skor"),
        (590, 210, "Review", "insan karari"),
        (850, 210, "Evaluation", "segment baseline"),
        (1110, 210, "Challenger", "offline + shadow"),
        (1370, 210, "Champion", "version + rollback"),
    ]
    for x, y, title, detail in nodes:
        draw.rounded_rectangle(
            (x, y, x + 210, y + 130), radius=18, fill="#E8EEF5", outline="#6B8CA5", width=3
        )
        draw.text((x + 16, y + 20), title, font=_font(21, True), fill="#17324D")
        draw.text((x + 16, y + 68), detail, font=_font(16), fill="#344054")
    for (x, y, _, _), (nx, ny, _, _) in zip(nodes, nodes[1:], strict=False):
        _arrow(draw, (x + 210, y + 65), (nx, ny + 65))
    draw.line([(1475, 340), (1475, 500), (435, 500), (435, 340)], fill="#26734D", width=4)
    draw.polygon([(435, 340), (428, 352), (442, 352)], fill="#26734D")
    draw.text(
        (750, 520),
        "Yalniz minimum veri ve kalite kapisi gecerse yeni model devreye girer",
        font=_font(19, True),
        fill="#26734D",
    )
    draw.text(
        (460, 600),
        "Risk/policy sinirlari model tarafindan degistirilemez",
        font=_font(19, True),
        fill="#9B1C1C",
    )
    img.save(path)


def set_run(
    run, size: float = 11, bold: bool = False, color: str = BLACK, italic: bool = False
) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(
    cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120
) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_last_image_alt(doc: Document, title: str, description: str) -> None:
    inline = doc.inline_shapes[-1]._inline
    inline.docPr.set("title", title)
    inline.docPr.set("descr", description)


def setup_docx() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(
        header.add_run("XNATIVE | MASTER ARCHITECTURE AND IMPLEMENTATION PLAN"), 8.5, True, MID_GRAY
    )
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(
        footer.add_run("Local-first | No X API | Human-approved public actions"),
        8.5,
        False,
        MID_GRAY,
    )
    return doc


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_paragraph(text, style=f"Heading {level}")


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), 11, True)
        set_run(p.add_run(text[len(bold_prefix) :]), 11)
    else:
        set_run(p.add_run(text), 11)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(text), 11)


def add_source_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 0.9
    set_run(p.add_run(text), 7)


def add_table(
    doc: Document, headers: list[str], rows: Iterable[Iterable[str]], widths_dxa: list[int]
) -> None:
    rows = [list(row) for row in rows]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        shade_cell(hdr[idx], LIGHT_GRAY)
        p = hdr[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(text), 9, True, NAVY)
    mark_header_row(table.rows[0])
    for data in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(str(text)), 8.5)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def build_docx(arch: Path, learn: Path) -> None:
    doc = setup_docx()
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("XNative Recursive Learning"), 28, True, NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Master Architecture and Implementation Plan"), 17, False, BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(
        p.add_run(
            "Local-first multimodal content intelligence | No mandatory X API or paid service"
        ),
        11,
        False,
        MID_GRAY,
        True,
    )
    doc.add_paragraph()
    add_table(
        doc,
        ["Field", "Value"],
        [
            (
                "Status",
                "Final implementation specification v1.0; application code remains phased work",
            ),
            ("Date", "2026-06-22"),
            (
                "Autonomous scope",
                "Capture processing, analysis, ranking, drafts, reports and controlled learning",
            ),
            ("Human-controlled scope", "Publishing, following and every public X action"),
        ],
        [2700, 6660],
    )
    doc.add_page_break()

    add_heading(doc, "1. Executive decision", 1)
    add_body(
        doc,
        "The supplied code is a useful prototype baseline, not a completed autonomous MVP. The binding implementation order begins with reliable capture and persistent data contracts, then adds multimodal intelligence and learning, and only then builds the operational review UI and final release documents.",
    )
    add_body(
        doc,
        "Autonomy boundary: internal analysis and learning may run without intervention; publishing, liking, following, reposting, replying, quoting or bookmarking on X remain user actions.",
        "Autonomy boundary:",
    )

    add_heading(doc, "2. Verified baseline and blockers", 1)
    verified = [
        "Fifteen baseline pytest tests pass in the pinned local environment.",
        "The ten-post fixture sample runs without an API key.",
        "No automatic public action exists in the reviewed code.",
        "Docker Compose configuration parses after creating .env.",
    ]
    blockers = [
        "The extension posts to /capture, but FastAPI does not register the route.",
        "The Streamlit script never calls run_app and has no review workflow.",
        "The sample pipeline does not persist data or consume learned weights.",
        "All fixture posts receive the same three template texts.",
        "The so-called pHash is a truncated exact SHA-256 and OCR is disconnected.",
        "Docker runtime, extension integration and end-to-end learning are unverified.",
    ]
    add_heading(doc, "Verified", 2)
    for item in verified:
        add_bullet(doc, item)
    add_heading(doc, "Release blockers", 2)
    for item in blockers:
        add_bullet(doc, item)

    add_heading(doc, "3. Target architecture", 1)
    add_body(
        doc,
        "The design uses a composable source -> hydration -> filtering -> feature -> scoring -> selection -> side-effect pipeline. This borrows the useful decomposition and diversity ideas from X's open-source For You architecture without copying its global corpus, distributed Kafka infrastructure or multi-gigabyte Phoenix serving assumptions.",
    )
    doc.add_picture(str(arch), width=Inches(6.35))
    set_last_image_alt(
        doc,
        "XNative hedef mimarisi",
        "Capture, local API, job queue, multimodal feature pipeline, event memory, ranking, review UI, storage and controlled learning flow.",
    )
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(
        caption.add_run("Figure 1. Local-first multimodal architecture and human action boundary."),
        9,
        False,
        MID_GRAY,
        True,
    )

    add_heading(doc, "4. Binding implementation order", 1)
    add_table(
        doc,
        ["Phase", "Deliverable", "Core change", "Acceptance gate"],
        PHASES,
        [650, 2350, 3180, 3180],
    )

    add_heading(doc, "5. Multimodal understanding strategy", 1)
    add_body(
        doc,
        "A post is not classified from explicit football words alone. Text, quote, OCR, image, video/audio, author history, event neighbors and cross-modal relationships remain separate evidence groups with confidence and missingness. The first production fusion is confidence-aware late fusion; a learned gated fusion model is a challenger only after labelled data is sufficient.",
    )
    for item in (
        "Relevant text plus unrelated media may be an intentional attention hook, not automatically spam.",
        "Low text-image similarity creates three hypotheses: wrong media, intentional contrast/irony, or model uncertainty.",
        "Event and author history decide between those hypotheses; user feedback updates format-family performance.",
        "Heavy VLM/LLM inference is an escalation path for uncertain high-value items, never a universal requirement.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "6. Retrieval, ranking and learning", 1)
    add_heading(doc, "Retrieval", 2)
    add_body(
        doc,
        "Start with SQLite FTS5 BM25 and multilingual dense embeddings combined by rank fusion. Add HNSW only after the corpus or latency benchmark exceeds the brute-force operating envelope.",
    )
    add_heading(doc, "Ranking", 2)
    add_body(
        doc,
        "Version zero is an explainable weighted utility model. It predicts or scores separate approve, edit, manual-post, engagement, ignore, reject and risk outcomes. A logistic or LightGBM ranker becomes a challenger after sufficient labels; a deep ranker is not justified by the initial data volume.",
    )
    add_heading(doc, "Controlled learning", 2)
    doc.add_picture(str(learn), width=Inches(6.35))
    set_last_image_alt(
        doc,
        "Kontrollu ogrenme dongusu",
        "Evidence, ranking, human review, segment evaluation, challenger testing, champion promotion and rollback loop.",
    )
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(
        caption.add_run("Figure 2. Feedback, evaluation, challenger and rollback loop."),
        9,
        False,
        MID_GRAY,
        True,
    )

    add_heading(doc, "7. Data lifecycle and archiving", 1)
    add_body(
        doc,
        "Link-only storage is space-efficient but loses evidence when a post is deleted or made private. Full-media archiving increases disk, copyright and privacy cost. The default is therefore metadata plus visible text/quote/alt text and provenance, with an optional thumbnail evidence layer and an explicit opt-in original-media layer.",
    )
    add_table(
        doc,
        ["Level", "Stored data", "Default"],
        [
            (
                "A - Metadata",
                "URL, visible text, quote, alt text, metrics, timestamps, provenance",
                "Yes",
            ),
            (
                "B - Minimal evidence",
                "Level A plus thumbnail/keyframe, OCR/ASR/caption and embedding",
                "Policy based",
            ),
            (
                "C - Original media",
                "Content-addressed original with license reason, quota and retention",
                "Opt-in",
            ),
        ],
        [1600, 5860, 1900],
    )
    add_body(
        doc,
        "The system never automates X bookmarks. It stores the local archive and opens the X link for a user-controlled bookmark or publish action.",
    )

    add_heading(doc, "8. Algorithm and complexity plan", 1)
    add_table(doc, ["Operation", "Method", "Time", "Memory"], COMPLEXITY, [2300, 3000, 2160, 1900])
    add_heading(doc, "Performance SLO draft", 2)
    add_table(doc, ["Surface", "Target", "Condition"], SLOS, [2450, 2050, 4860])
    add_body(
        doc,
        "All optimization is benchmark-driven. Cache keys include content hash, model version and configuration hash. Model inference is lazy and batched; UI queries are paginated; video frames and media size are strictly bounded.",
    )

    add_heading(doc, "9. Algorithm selection and promotion", 1)
    add_body(
        doc,
        "The project uses a baseline -> challenger -> shadow -> champion ladder. Model names do not justify adoption: each challenger must improve time-split quality and calibration without violating risk recall, p95 latency, RAM/VRAM or rollback gates.",
    )
    add_table(
        doc,
        ["Problem", "Default", "Challenger", "Explicit guardrail"],
        ALGORITHM_LADDER,
        [1100, 2550, 2550, 3160],
    )
    add_body(
        doc,
        "Data policy: below 200 clean labels use rules/retrieval and speed baselines; from 200 use regularized logistic/online SGD; from 2,000 test nonlinear tabular fusion; pairwise ranking requires at least 20,000 preference pairs. LSTM/GRU remains a challenger only after at least 50,000 clean sequence labels and a learning-curve case. These are initial operating gates, not universal constants.",
    )
    for item in (
        "Vision: benchmark a small pretrained OpenCLIP/MobileCLIP class for image-text evidence; do not train CLIP locally or treat it as OCR/fact verification.",
        "Graph: build typed SQLite graph plus metapath and Personalized PageRank baselines before Node2Vec, R-GCN, GraphSAGE or HGT.",
        "Vectors: exact ground-truth -> FAISS IndexFlatIP -> HNSW/IVF-PQ only at corpus/SLO thresholds. Milvus is rejected on the 8 GB target host.",
        "Pipeline: transactional outbox, priority aging, bounded backpressure, micro-batch, full-jitter retry and token bucket precede distributed orchestration.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "10. Target deployment hardware", 1)
    add_body(
        doc,
        "The always-on computer is a ninth-generation Intel Core i5 with 8 GB RAM and a GeForce GTX 1050. The GTX 1050 VRAM amount must be detected during installation; scheduling assumes a conservative 2 GB floor and supports full CPU-only operation.",
    )
    add_table(
        doc,
        ["Resource", "Operating rule"],
        [
            ("RAM", "Reserve at least 2 GB for OS/browser; application soft limit 5.5 GB"),
            ("Concurrency", "One heavy model worker and at most two light I/O/CPU workers"),
            ("OCR", "Tesseract on CPU; EasyOCR only after measured quality uplift"),
            ("Text", "384-dimensional multilingual ONNX/int8 embedding"),
            ("Vision", "Small image encoder, batch 1-4; GPU OOM falls back to CPU"),
            ("Audio/video", "Whisper tiny/base int8 and 3-8 sampled frames only when needed"),
            ("Training", "Bounded online weights and classical ML; no local deep-model training"),
        ],
        [2200, 7160],
    )

    add_heading(doc, "11. Optional provider and Google AI policy", 1)
    add_body(
        doc,
        "The core pipeline must remain operational without keys or quotas. External providers may be used only behind adapters, with explicit enablement, data minimization, timeout, 429 handling, circuit breakers and a local fallback. Free quotas are re-verified from official sources at activation time and are never treated as a service-level guarantee.",
    )
    for item in (
        "Groq may serve as optional text or ASR inference under published free-plan rate limits.",
        "Hugging Face models may be pinned and cached locally; each model license is reviewed separately.",
        "Google AI Pro may include higher AI Studio limits and an eligible monthly USD 10 Cloud credit, but the benefit must be redeemed and verified on the billing account.",
        "The 1,000 Google AI Pro AI credits are for Flow/Whisk, not Gemini API tokens.",
        "Gemini API Free/Paid tier, RPM/TPM/RPD and costs remain project- and model-specific; the API key page and AI Studio billing view are authoritative.",
        "Gemini Flash-Lite/Flash is the first optional multimodal escalation; Pro is reserved for low-volume difficult cases.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "12. Product and visualization brief", 1)
    add_body(
        doc,
        "The operational workspace requires Dashboard, Capture Inbox, Suggestions Queue, Source Candidates, Style Memory, Reports and Settings. It is table/list driven, shows essential evidence without hover, exposes live/stale/offline/partial states, and preserves keyboard and mobile review paths.",
    )
    add_body(
        doc,
        "A visual target must be approved before frontend implementation. The brief is followed by exactly three full-screen and mobile concept directions; the chosen concept becomes the implementation contract.",
    )

    add_heading(doc, "13. Release evidence contract", 1)
    for item in (
        "Every code change references a backlog ID.",
        "Every completed requirement has code/migration, automated test, performance/security impact and document evidence.",
        "A model becomes champion only after offline and shadow gates; rollback remains available.",
        "Final Markdown, DOCX and PDF carry the same release or commit identity.",
        "No Final QA statement is inferred from file presence or unit tests alone.",
    ):
        add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "14. Implementation-ready contract", 1)
    add_body(
        doc,
        "The documentation set is final for implementation version 1.0. This means architecture and coding decisions are closed; it does not claim the application is release-ready. Code completion remains governed by the backlog and traceability evidence.",
    )
    add_table(
        doc,
        ["Authority", "Contract"],
        [
            ("MASTER_IMPLEMENTATION_BACKLOG", "Phase order, dependencies and acceptance gates"),
            (
                "IMPLEMENTATION_SPECIFICATION",
                "Domain, DB, API, job, config, errors and compatibility",
            ),
            ("TEST_PLAN", "Named phase acceptance and release evidence"),
            ("ALGORITHM_TECHNOLOGY_RADAR", "Use, benchmark, threshold and reject decisions"),
            ("RISK / RUNBOOK", "Operational mitigation, incident, backup and rollback"),
        ],
        [3100, 6260],
    )
    for item in (
        "API contract: /api/v1 is canonical; POST /capture is a temporary compatibility alias; capture commits inbox and job before returning 202.",
        "Queue contract: pending/running/retry/completed/dead with leases, full-jitter retry, priority aging and heavy=1/light=2 resource limits.",
        "Data contract: UUIDv4, UTC timestamps, migration checksums, append-only feedback/audit and model-feature-config provenance.",
        "Phase 1 starts with domain contracts, SQL migrations, repositories, WAL/idempotency and audit. ML, FAISS, graph models and UI cannot bypass it.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "15. Sources", 1)
    for source in SOURCES:
        add_source_bullet(doc, source)

    doc.core_properties.title = "XNative Master Architecture and Implementation Plan"
    doc.core_properties.subject = "Local-first multimodal X content intelligence"
    doc.core_properties.author = "XNative Project"
    doc.core_properties.keywords = "XNative, multimodal, NLP, CV, ranking, local-first"
    doc.save(DOCX_OUT)


def pdf_styles():
    pdfmetrics.registerFont(TTFont("Arial", "/System/Library/Fonts/Supplemental/Arial.ttf"))
    pdfmetrics.registerFont(
        TTFont("Arial-Bold", "/System/Library/Fonts/Supplemental/Arial Bold.ttf")
    )
    pdfmetrics.registerFont(
        TTFont("Arial-Italic", "/System/Library/Fonts/Supplemental/Arial Italic.ttf")
    )
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="XTitle",
            fontName="Arial-Bold",
            fontSize=25,
            leading=30,
            textColor=colors.HexColor("#17324D"),
            alignment=TA_CENTER,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="XSubtitle",
            fontName="Arial",
            fontSize=14,
            leading=18,
            textColor=colors.HexColor("#2E74B5"),
            alignment=TA_CENTER,
            spaceAfter=14,
        )
    )
    styles.add(
        ParagraphStyle(
            name="XH1",
            fontName="Arial-Bold",
            fontSize=16,
            leading=20,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=12,
            spaceAfter=8,
            keepWithNext=True,
        )
    )
    styles.add(
        ParagraphStyle(
            name="XH2",
            fontName="Arial-Bold",
            fontSize=12.5,
            leading=16,
            textColor=colors.HexColor("#17324D"),
            spaceBefore=9,
            spaceAfter=5,
            keepWithNext=True,
        )
    )
    styles.add(
        ParagraphStyle(
            name="XBody",
            fontName="Arial",
            fontSize=9.5,
            leading=13,
            textColor=colors.HexColor("#111111"),
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="XBullet",
            fontName="Arial",
            fontSize=9.2,
            leading=12.5,
            leftIndent=14,
            firstLineIndent=-8,
            bulletIndent=6,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="XCaption",
            fontName="Arial-Italic",
            fontSize=8,
            leading=10,
            textColor=colors.HexColor("#667085"),
            alignment=TA_CENTER,
            spaceAfter=8,
        )
    )
    return styles


def p(text: str, style) -> Paragraph:
    return Paragraph(text.replace("&", "&amp;"), style)


def pdf_table(
    headers: list[str], rows: Iterable[Iterable[str]], widths: list[float], styles
) -> Table:
    data = [
        [
            p(
                h,
                ParagraphStyle(
                    name=f"th{idx}{id(headers)}",
                    parent=styles["XBody"],
                    fontName="Arial-Bold",
                    fontSize=7.3,
                    leading=9,
                    textColor=colors.HexColor("#17324D"),
                ),
            )
            for idx, h in enumerate(headers)
        ]
    ]
    for row in rows:
        data.append(
            [
                p(
                    str(value),
                    ParagraphStyle(
                        name=f"td{idx}{len(data)}{id(row)}",
                        parent=styles["XBody"],
                        fontSize=7.1,
                        leading=9,
                    ),
                )
                for idx, value in enumerate(row)
            ]
        )
    table = Table(data, colWidths=[width * inch for width in widths], repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C8D0D8")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def add_page_number(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Arial", 8)
    canvas.setFillColor(colors.HexColor("#667085"))
    canvas.drawString(inch, 0.48 * inch, "XNATIVE | MASTER ARCHITECTURE AND IMPLEMENTATION PLAN")
    canvas.drawRightString(7.5 * inch, 0.48 * inch, f"Page {doc.page}")
    canvas.restoreState()


def build_pdf(arch: Path, learn: Path) -> None:
    styles = pdf_styles()
    story: list[Flowable] = []
    story.extend(
        [
            Spacer(1, 1.1 * inch),
            p("XNative Recursive Learning", styles["XTitle"]),
            p("Master Architecture and Implementation Plan", styles["XSubtitle"]),
            p(
                "Local-first multimodal content intelligence | No mandatory X API or paid service",
                ParagraphStyle(
                    name="cover",
                    parent=styles["XBody"],
                    alignment=TA_CENTER,
                    textColor=colors.HexColor("#667085"),
                    fontName="Arial-Italic",
                ),
            ),
            Spacer(1, 0.35 * inch),
        ]
    )
    story.append(
        pdf_table(
            ["Field", "Value"],
            [
                (
                    "Status",
                    "Final implementation specification v1.0; code remains phased work",
                ),
                ("Date", "2026-06-22"),
                ("Autonomous scope", "Analysis, ranking, drafts, reports and controlled learning"),
                ("Human-controlled scope", "Publishing, following and every public X action"),
            ],
            [1.55, 4.95],
            styles,
        )
    )
    story.append(PageBreak())

    def h1(text):
        story.append(p(text, styles["XH1"]))

    def h2(text):
        story.append(p(text, styles["XH2"]))

    def body(text):
        story.append(p(text, styles["XBody"]))

    def bullet(text):
        story.append(Paragraph(text.replace("&", "&amp;"), styles["XBullet"], bulletText="-"))

    h1("1. Executive decision")
    body(
        "The supplied code is a useful prototype baseline, not a completed autonomous MVP. Implementation begins with reliable capture and persistent contracts, then multimodal intelligence and learning, and only then the review UI and release documents."
    )
    body(
        "Autonomy covers internal processing. Publishing, liking, following, reposting, replying, quoting or bookmarking on X remain user actions."
    )
    h1("2. Verified baseline and blockers")
    h2("Verified")
    for item in (
        "Fifteen baseline pytest tests pass in the pinned local environment.",
        "The ten-post fixture sample runs without an API key.",
        "No automatic public action exists in the reviewed code.",
        "Docker Compose configuration parses after creating .env.",
    ):
        bullet(item)
    h2("Release blockers")
    for item in (
        "The extension posts to /capture, but FastAPI does not register the route.",
        "The Streamlit script never calls run_app and has no review workflow.",
        "The sample pipeline does not persist data or consume learned weights.",
        "All fixture posts receive the same three template texts.",
        "The current pHash is exact-byte based and OCR is disconnected.",
        "Docker runtime and end-to-end learning are unverified.",
    ):
        bullet(item)
    h1("3. Target architecture")
    body(
        "The design uses a composable source -> hydration -> filtering -> feature -> scoring -> selection -> side-effect pipeline. It adapts X's public recommendation decomposition without adopting global corpus, Kafka or multi-gigabyte serving assumptions."
    )
    story.append(RLImage(str(arch), width=6.5 * inch, height=2.74 * inch))
    story.append(
        p(
            "Figure 1. Local-first multimodal architecture and human action boundary.",
            styles["XCaption"],
        )
    )
    h1("4. Binding implementation order")
    story.append(
        pdf_table(
            ["Phase", "Deliverable", "Core change", "Acceptance gate"],
            PHASES,
            [0.45, 1.65, 2.25, 2.15],
            styles,
        )
    )
    h1("5. Multimodal understanding strategy")
    body(
        "A post is not classified from explicit football words alone. Text, quote, OCR, image, video/audio, author history, event neighbors and cross-modal relationships remain separate evidence groups with confidence and missingness."
    )
    for item in (
        "Relevant text plus unrelated media may be an intentional attention hook, not automatically spam.",
        "Low text-image similarity creates hypotheses for wrong media, intentional contrast/irony or model uncertainty.",
        "Event and author history decide between hypotheses; feedback updates format-family performance.",
        "Heavy VLM/LLM inference is reserved for uncertain high-value items.",
    ):
        bullet(item)
    h1("6. Retrieval, ranking and controlled learning")
    body(
        "Retrieval begins with FTS5 BM25 plus multilingual embeddings. Ranking starts as explainable multi-action utility; logistic or LightGBM models are challengers after sufficient labels. Contextual bandits explore only safe tone/format choices."
    )
    story.append(RLImage(str(learn), width=6.5 * inch, height=2.6 * inch))
    story.append(
        p("Figure 2. Feedback, evaluation, challenger and rollback loop.", styles["XCaption"])
    )
    h1("7. Data lifecycle and archiving")
    body(
        "Link-only storage can lose evidence, while full-media archiving adds disk, copyright and privacy cost. The default is metadata and visible evidence, with optional thumbnails and explicit opt-in original media."
    )
    story.append(
        pdf_table(
            ["Level", "Stored data", "Default"],
            [
                (
                    "A - Metadata",
                    "URL, visible text, quote, alt text, metrics, timestamps, provenance",
                    "Yes",
                ),
                (
                    "B - Minimal evidence",
                    "Thumbnail/keyframe, OCR/ASR/caption and embedding",
                    "Policy based",
                ),
                (
                    "C - Original media",
                    "Content-addressed original with reason, quota and retention",
                    "Opt-in",
                ),
            ],
            [1.1, 4.25, 1.15],
            styles,
        )
    )
    body(
        "The system never automates X bookmarks. Local archive and X links remain separate; X actions stay user-controlled."
    )
    h1("8. Algorithm and complexity plan")
    story.append(
        pdf_table(
            ["Operation", "Method", "Time", "Memory"], COMPLEXITY, [1.45, 2.2, 1.55, 1.3], styles
        )
    )
    h2("Performance SLO draft")
    story.append(pdf_table(["Surface", "Target", "Condition"], SLOS, [1.6, 1.45, 3.45], styles))
    body(
        "Optimization is benchmark-driven. Cache keys include content hash, model version and config hash. Models are lazy/batched, UI queries paginated, and media/video strictly bounded."
    )
    h1("9. Algorithm selection and promotion")
    body(
        "The project uses a baseline -> challenger -> shadow -> champion ladder. A challenger must improve time-split quality and calibration without violating risk recall, p95 latency, RAM/VRAM or rollback gates."
    )
    story.append(
        pdf_table(
            ["Problem", "Default", "Challenger", "Guardrail"],
            ALGORITHM_LADDER,
            [0.8, 1.8, 1.8, 2.1],
            styles,
        )
    )
    body(
        "Below 200 clean labels use rules/retrieval and speed baselines; from 200 use regularized logistic/online SGD; from 2,000 test nonlinear fusion; pairwise ranking requires 20,000 preference pairs. LSTM/GRU requires at least 50,000 clean sequence labels and measured learning-curve uplift."
    )
    for item in (
        "Use a small pretrained CLIP class as a benchmark, not as OCR or fact verification.",
        "Build typed graph plus metapath/PPR before Node2Vec or heterogeneous GNNs.",
        "Use exact vectors, then FAISS at the measured threshold; Milvus is rejected on 8 GB RAM.",
        "Use bounded local queue algorithms before distributed orchestration.",
    ):
        bullet(item)
    h1("10. Target deployment hardware")
    body(
        "The always-on machine is a ninth-generation Intel Core i5 with 8 GB RAM and a GeForce GTX 1050. Installation detects VRAM and driver compatibility; the default scheduler assumes a 2 GB VRAM floor and preserves CPU-only operation."
    )
    story.append(
        pdf_table(
            ["Resource", "Operating rule"],
            [
                ("RAM", "2 GB OS/browser reserve; 5.5 GB application soft limit"),
                ("Concurrency", "One heavy worker plus at most two light workers"),
                ("Models", "Small quantized text/vision; only one heavy model loaded"),
                ("OCR/ASR", "Tesseract CPU; Whisper tiny/base int8 on selected media"),
                ("Training", "Classical ML only; no local deep-model training"),
            ],
            [1.6, 4.9],
            styles,
        )
    )
    h1("11. Optional provider and Google AI policy")
    body(
        "The core pipeline remains usable without keys. Providers are optional adapters with explicit enablement, data minimization, timeout, 429 handling, circuit breakers and local fallback. Quotas are re-verified from official sources when activated."
    )
    for item in (
        "Google AI Pro can provide higher AI Studio limits and may expose an eligible monthly USD 10 Cloud credit after account redemption.",
        "The plan's 1,000 AI credits are Flow/Whisk credits, not Gemini API token balance.",
        "Gemini API tier, active limits and billing are separate project-level settings.",
        "Flash-Lite/Flash is the default escalation; Pro is reserved for difficult low-volume cases.",
    ):
        bullet(item)
    h1("12. Product brief")
    body(
        "The workspace includes Dashboard, Capture Inbox, Suggestions Queue, Source Candidates, Style Memory, Reports and Settings. Essential evidence is visible without hover; live/stale/offline/partial states, keyboard access and mobile review are required."
    )
    body(
        "Frontend implementation waits for brief approval and selection among three full-screen/mobile visual directions."
    )
    h1("13. Release evidence contract")
    for item in (
        "Every code change references a backlog ID.",
        "Completion requires code, automated test, performance/security impact and document evidence.",
        "Champion promotion requires offline and shadow gates plus rollback.",
        "Markdown, DOCX and PDF describe the same release.",
        "Final QA never infers completion from file presence or unit tests alone.",
    ):
        bullet(item)
    story.append(PageBreak())
    h1("14. Implementation-ready contract")
    body(
        "Documentation version 1.0 closes architecture and coding decisions without claiming the application is release-ready. Backlog and traceability evidence remain the completion authority."
    )
    story.append(
        pdf_table(
            ["Authority", "Contract"],
            [
                ("Backlog", "Phase order, dependencies and acceptance gates"),
                ("Implementation specification", "Domain, DB, API, jobs, config and errors"),
                ("Test plan", "Named acceptance IDs and release evidence"),
                ("Algorithm radar", "Use, benchmark, threshold and reject decisions"),
                ("Risk / runbook", "Incident, backup, fallback and rollback"),
            ],
            [2.1, 4.4],
            styles,
        )
    )
    for item in (
        "/api/v1 is canonical; /capture is a temporary alias; 202 follows inbox+job commit.",
        "Durable jobs use leases, jittered retry, priority aging and bounded resource classes.",
        "Data uses UUIDv4, UTC, migration checksums and complete provenance.",
        "Phase 1 domain/storage work precedes models, vector indexes, graph ML and UI.",
    ):
        bullet(item)
    h1("15. Sources")
    for source in SOURCES:
        bullet(source)

    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=LETTER,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=0.85 * inch,
        bottomMargin=0.75 * inch,
        title="XNative Master Architecture and Implementation Plan",
        author="XNative Project",
    )
    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    TMP.mkdir(parents=True, exist_ok=True)
    arch = TMP / "architecture_overview.png"
    learn = TMP / "learning_loop.png"
    create_architecture_image(arch)
    create_learning_image(learn)
    build_docx(arch, learn)
    build_pdf(arch, learn)
    print(DOCX_OUT)
    print(PDF_OUT)


if __name__ == "__main__":
    main()
